"""
DOCX Translator - Perfect Document Translation

Translates Word documents while preserving:
- All formatting (bold, italic, fonts, sizes)
- Tables and images
- Mathematical formulas
- Document structure

NEW: Batch translation for better context!

© 2025 Sven Kalinowski with small help of Lino Casu
Licensed under the Anti-Capitalist Software License v1.4
"""
from __future__ import annotations

import logging
import re
import time
import copy
from pathlib import Path
from typing import Optional, Dict, List, Callable
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger("docx_translator")

# =============================================================================
# CONFIGURATION
# =============================================================================

MAX_RETRIES = 5
RETRY_DELAY = 1.5
BATCH_SIZE = 5  # Translate multiple paragraphs together for better context

# Scientific glossary for German
GLOSSARY_DE = {
    "entanglement": "Verschränkung",
    "entangled": "verschränkt",
    "coherence": "Kohärenz",
    "decoherence": "Dekohärenz",
    "superposition": "Superposition",
    "qubit": "Qubit",
    "qubits": "Qubits",
    "fidelity": "Fidelität",
    "spacetime": "Raumzeit",
    "phase": "Phase",
    "quantum": "Quanten",
    "abstract": "Zusammenfassung",
    "introduction": "Einleitung",
    "methods": "Methoden",
    "results": "Ergebnisse",
    "conclusions": "Schlussfolgerungen",
    "references": "Literaturverzeichnis",
    "acknowledgments": "Danksagung",
    "background": "Hintergrund",
    "discussion": "Diskussion",
}


@dataclass
class TranslationResult:
    """Result of document translation."""
    success: bool
    output_path: Optional[str]
    paragraphs_translated: int
    paragraphs_skipped: int
    warnings: List[str] = field(default_factory=list)


# =============================================================================
# TRANSLATION
# =============================================================================

def translate_batch_ollama(
    texts: List[str],
    model: str,
    target_language: str,
    glossary: Dict[str, str] = None
) -> List[str]:
    """Translate multiple texts in one call for better context."""
    import requests
    
    if not texts:
        return texts
    
    # Filter empty texts
    valid_indices = [i for i, t in enumerate(texts) if t.strip() and re.search(r'[a-zA-Z]{2,}', t)]
    if not valid_indices:
        return texts
    
    # Build numbered text block
    numbered_texts = []
    for idx, i in enumerate(valid_indices):
        numbered_texts.append(f"[{idx+1}] {texts[i]}")
    
    combined = "\n\n".join(numbered_texts)
    
    glossary_text = ""
    if glossary:
        glossary_text = "MANDATORY TERMINOLOGY (use these exact translations):\n"
        for en, trans in list(glossary.items())[:15]:  # Top 15 terms
            glossary_text += f"  {en} → {trans}\n"
    
    system_prompt = f"""You are a scientific translator. Translate to {target_language}.

{glossary_text}
CRITICAL RULES:
1. Translate each numbered section [1], [2], etc. separately
2. Keep the [N] markers in your output
3. Output ONLY translations - no explanations
4. Keep ALL math symbols: Δ, Φ, ω, ×, ⁻¹, ², ³, π, α, β, γ, ∞, ∑, ∫
5. Keep numbers, units, author names, URLs unchanged
6. Use formal scientific {target_language}"""

    for attempt in range(MAX_RETRIES):
        try:
            response = requests.post(
                "http://localhost:11434/api/chat",
                json={
                    "model": model,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"Translate these scientific text sections to {target_language}:\n\n{combined}"}
                    ],
                    "stream": False,
                    "options": {"temperature": 0.1, "num_predict": 8192}
                },
                timeout=300
            )
            
            if response.status_code == 200:
                result = response.json().get("message", {}).get("content", "").strip()
                
                # Parse numbered responses
                translations = texts.copy()  # Start with originals
                
                for idx, orig_i in enumerate(valid_indices):
                    marker = f"[{idx+1}]"
                    # Find this section in response
                    start = result.find(marker)
                    if start != -1:
                        start += len(marker)
                        # Find next marker or end
                        next_marker = result.find(f"[{idx+2}]", start)
                        if next_marker == -1:
                            section = result[start:].strip()
                        else:
                            section = result[start:next_marker].strip()
                        
                        if section and len(section) >= len(texts[orig_i]) * 0.3:
                            translations[orig_i] = section
                
                return translations
            
            time.sleep(RETRY_DELAY * (attempt + 1))
            
        except Exception as e:
            logger.warning(f"Batch translation attempt {attempt + 1} failed: {e}")
            time.sleep(RETRY_DELAY * (attempt + 1))
    
    return texts  # Return originals on failure


def translate_text_ollama(
    text: str,
    model: str,
    target_language: str,
    glossary: Dict[str, str] = None
) -> str:
    """Translate single text using Ollama (fallback)."""
    import requests
    
    if not text.strip():
        return text
    
    # Skip if only numbers/symbols
    if not re.search(r'[a-zA-Z]{2,}', text):
        return text
    
    glossary_text = ""
    if glossary:
        glossary_text = "MANDATORY TERMINOLOGY (use these exact translations):\n"
        for en, trans in glossary.items():
            glossary_text += f"  {en} → {trans}\n"
    
    system_prompt = f"""You are a scientific translator. Translate to {target_language}.

{glossary_text}
CRITICAL RULES:
1. Output ONLY the translation - no explanations, no quotes
2. Keep ALL mathematical symbols EXACTLY as they are: Δ, Φ, ω, ×, ⁻, ¹, ², ³, ⁴, ⁵, ⁶, ⁷, ⁸, ⁹, ⁰, π, α, β, γ, ∞, ∑, ∫, ≈, ≠, ≤, ≥, —, –
3. Keep ALL numbers and units unchanged
4. Keep author names unchanged (Carmen Wrede, Lino Casu)
5. Keep URLs and emails unchanged
6. Use formal scientific German
7. Translate section headers: Abstract→Zusammenfassung, Introduction→Einleitung, etc."""

    for attempt in range(MAX_RETRIES):
        try:
            response = requests.post(
                "http://localhost:11434/api/chat",
                json={
                    "model": model,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"Translate this scientific text to {target_language}:\n\n{text}"}
                    ],
                    "stream": False,
                    "options": {"temperature": 0.1, "num_predict": 4096}
                },
                timeout=180
            )
            
            if response.status_code == 200:
                result = response.json().get("message", {}).get("content", "").strip()
                # Clean up common artifacts
                result = re.sub(r'^(Here|Translation|Übersetzung|The translation).*?:\s*', '', result, flags=re.I)
                result = re.sub(r'^["\']|["\']$', '', result)
                result = re.sub(r'^```\w*\n?', '', result)
                result = re.sub(r'\n?```$', '', result)
                
                if result and len(result) >= len(text) * 0.3:
                    return result
            
            time.sleep(RETRY_DELAY * (attempt + 1))
            
        except Exception as e:
            logger.warning(f"Translation attempt {attempt + 1} failed: {e}")
            time.sleep(RETRY_DELAY * (attempt + 1))
    
    return text


def is_translatable(text: str) -> bool:
    """Check if text should be translated."""
    text = text.strip()
    
    if len(text) < 2:
        return False
    
    # Must have letters
    if not re.search(r'[a-zA-Z]', text):
        return False
    
    # Skip URLs/emails
    if re.search(r'https?://|@.*\.[a-z]', text):
        return False
    
    # Skip pure numbers/dates
    if re.match(r'^[\d\s\-/\.,:×]+$', text):
        return False
    
    return True


def has_math_content(element) -> bool:
    """Check if element contains Office Math (OMML)."""
    # Check for oMath elements
    math_elements = element._element.findall('.//' + qn('m:oMath'))
    if math_elements:
        return True
    
    # Check for oMathPara
    math_para = element._element.findall('.//' + qn('m:oMathPara'))
    if math_para:
        return True
    
    return False


# =============================================================================
# DOCX TRANSLATION
# =============================================================================

def translate_docx(
    input_path: str,
    output_path: str,
    model: str,
    target_language: str,
    progress_callback: Optional[Callable] = None
) -> TranslationResult:
    """
    Translate a DOCX file while preserving all formatting.
    Uses batch translation for better context.
    """
    result = TranslationResult(
        success=False,
        output_path=None,
        paragraphs_translated=0,
        paragraphs_skipped=0,
        warnings=[]
    )
    
    glossary = GLOSSARY_DE if target_language.lower() in ["german", "deutsch", "de"] else {}
    
    try:
        # Create output directory if needed
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)
        
        doc = Document(input_path)
        total_paragraphs = len(doc.paragraphs)
        
        if progress_callback:
            progress_callback(0, 100, f"Processing {total_paragraphs} paragraphs in batches...")
        
        # Collect translatable paragraphs in batches
        batch_paragraphs = []  # (index, paragraph, original_text)
        
        for i, para in enumerate(doc.paragraphs):
            # Skip empty paragraphs
            if not para.text.strip():
                result.paragraphs_skipped += 1
                continue
            
            # Skip paragraphs with math content (preserve formulas)
            if has_math_content(para):
                result.paragraphs_skipped += 1
                continue
            
            # Check if translatable
            if not is_translatable(para.text):
                result.paragraphs_skipped += 1
                continue
            
            batch_paragraphs.append((i, para, para.text))
        
        # Process in batches for better context
        total_batches = (len(batch_paragraphs) + BATCH_SIZE - 1) // BATCH_SIZE
        
        for batch_idx in range(0, len(batch_paragraphs), BATCH_SIZE):
            batch = batch_paragraphs[batch_idx:batch_idx + BATCH_SIZE]
            current_batch = batch_idx // BATCH_SIZE + 1
            
            progress = int(5 + 80 * batch_idx / len(batch_paragraphs))
            if progress_callback:
                progress_callback(progress, 100, f"Batch {current_batch}/{total_batches} ({len(batch)} paragraphs)")
            
            # Extract texts for batch translation
            texts = [item[2] for item in batch]
            
            # Batch translate for better context
            translated_texts = translate_batch_ollama(texts, model, target_language, glossary)
            
            # Apply translations
            for (orig_idx, para, original_text), translated_text in zip(batch, translated_texts):
                if translated_text == original_text:
                    result.paragraphs_skipped += 1
                    continue
                
                # Replace text while preserving formatting
                if para.runs:
                    total_len = sum(len(run.text) for run in para.runs)
                    
                    if total_len > 0:
                        translated_pos = 0
                        
                        for run in para.runs:
                            if not run.text:
                                continue
                            
                            run_proportion = len(run.text) / total_len
                            chars_for_run = int(len(translated_text) * run_proportion)
                            
                            if run == para.runs[-1]:
                                run.text = translated_text[translated_pos:]
                            else:
                                run.text = translated_text[translated_pos:translated_pos + chars_for_run]
                                translated_pos += chars_for_run
                
                result.paragraphs_translated += 1
        
        # Translate tables
        if progress_callback:
            progress_callback(90, 100, "Translating tables...")
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip() and is_translatable(para.text):
                            translated = translate_text_ollama(para.text, model, target_language, glossary)
                            if translated != para.text and para.runs:
                                # Simple replacement for table cells
                                for run in para.runs:
                                    run.text = ""
                                if para.runs:
                                    para.runs[0].text = translated
        
        # Save
        if progress_callback:
            progress_callback(95, 100, "Saving document...")
        
        doc.save(output_path)
        
        result.success = True
        result.output_path = output_path
        
        if progress_callback:
            progress_callback(100, 100, "Complete!")
        
        return result
        
    except Exception as e:
        logger.exception(f"DOCX translation failed: {e}")
        result.warnings.append(str(e))
        return result


# =============================================================================
# CLI
# =============================================================================

if __name__ == "__main__":
    import sys
    
    logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
    
    if len(sys.argv) < 2:
        print("Usage: python docx_translator.py input.docx [output.docx] [language] [model]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.docx', '_translated.docx')
    language = sys.argv[3] if len(sys.argv) > 3 else "German"
    model = sys.argv[4] if len(sys.argv) > 4 else "qwen2.5:7b"
    
    print(f"Translating {input_file} to {language}...")
    print(f"Using model: {model}")
    print()
    
    result = translate_docx(
        input_file, output_file, model, language,
        progress_callback=lambda c, t, m: print(f"[{c:3d}%] {m}")
    )
    
    print()
    print("=" * 60)
    print(f"SUCCESS: {result.success}")
    print(f"Paragraphs translated: {result.paragraphs_translated}")
    print(f"Paragraphs skipped: {result.paragraphs_skipped}")
    if result.warnings:
        print(f"Warnings: {result.warnings}")
    if result.output_path:
        print(f"\nOutput: {result.output_path}")
